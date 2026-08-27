"""Word figure pack export for report drafters."""

from __future__ import annotations

from io import BytesIO
from typing import Mapping


def build_figure_docx_bytes(
    *,
    png_bytes: bytes,
    caption: str,
    title: str,
    metadata: Mapping[str, object] | None = None,
) -> bytes:
    """Embed prepared PNG with caption and optional metadata table."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - optional dep
        raise RuntimeError(
            "python-docx is required for Word export. Install with: pip install python-docx"
        ) from exc

    document = Document()
    document.add_heading(title or "Cross Section", level=1)
    if caption:
        document.add_paragraph(caption)
    if png_bytes:
        stream = BytesIO(png_bytes)
        document.add_picture(stream, width=Inches(6.5))
    if metadata:
        document.add_heading("Figure metadata", level=2)
        table = document.add_table(rows=1, cols=2)
        header = table.rows[0].cells
        header[0].text = "Field"
        header[1].text = "Value"
        for key, value in metadata.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
